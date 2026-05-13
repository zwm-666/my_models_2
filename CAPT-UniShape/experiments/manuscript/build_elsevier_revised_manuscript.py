"""Build the revised Elsevier-style CAPT-UniShape manuscript DOCX.

The manuscript content is deliberately limited to the supplied and locally
verified project evidence: the seed-44 fixed 8:2 protocol, the ablation table,
and the joint-input noise experiments.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "CAPT-UniShape_Elsevier_revised_draft.docx"
MD_PATH = OUT_DIR / "CAPT-UniShape_Elsevier_revised_draft.md"


TITLE = "CAPT-UniShape: Condition-Aware Prototype Transport UniShape for PEMFC Fault Classification"

HIGHLIGHTS = [
    "CAPT-UniShape combines stack operation sequences, EIS/impedance shape sequences, and condition/statistical vectors for PEMFC fault classification.",
    "Residual KAN-Fusion uses a stable MLP branch plus a bottleneck KAN residual branch to model nonlinear cross-modal interactions.",
    "A condition-aware RBF prototype transport head adapts class prototypes from the encoded condition token before cosine classification.",
    "Under the evaluated seed-44 fixed 8:2 test protocol, the model achieved 100.00% accuracy and 100.00% Macro-F1 on 55 test windows.",
    "Noise experiments show that joint input perturbations reduce Macro-F1 and expose sensitivity of the normal class, so strong-noise robustness remains an open issue.",
]

ABSTRACT = (
    "Reliable fault classification is important for proton exchange membrane fuel cell (PEMFC) systems because "
    "water management faults, reactant starvation, and condition-dependent operating drift can reduce efficiency and "
    "accelerate degradation. Existing data-driven diagnostic models often rely on a single signal source or a static "
    "classification boundary, which may be insufficient when operational time series, impedance-related responses, "
    "and operating conditions jointly shape the observed fault signature. This paper presents CAPT-UniShape, a "
    "Condition-Aware Prototype Transport UniShape model for PEMFC fault classification. The model receives three "
    "inputs: stack operation sequences x_op in R^{N x 3 x 64}, constructed EIS/impedance statistical shape sequences "
    "x_eis in R^{N x 4 x 128}, "
    "and condition/statistical vectors x_cond in R^{N x 12}. Two official UniShape backbones extract operation and "
    "EIS shape embeddings, while a condition encoder maps x_cond into a condition token. The concatenated embeddings "
    "are fused by Residual KAN-Fusion, which combines an MLP branch with a bottleneck KAN residual branch. The final "
    "classifier uses a condition-aware RBF prototype transport head, where each static class prototype is shifted by "
    "a condition-dependent offset before cosine-logit classification. Under the evaluated seed-44 fixed 8:2 test "
    "protocol, Official-CAPT-UniShape-RBF-KANFusion achieved 100.00% test accuracy, 100.00% test Macro-F1, and "
    "100.00% test Weighted-F1 on 55 test windows, with 6,489,373 trainable parameters. Ablation results indicate that the EIS/impedance branch is critical, "
    "and that RBF prototypes and Residual KAN-Fusion improve the normal-class recall in the evaluated split. However, "
    "static prototypes and removal of transport or separation regularization also reached 100.00%, so their marginal "
    "contribution should not be overstated. Joint-input noise experiments further show marked Macro-F1 degradation "
    "and normal-class sensitivity, indicating that robust training and broader validation are needed before drawing "
    "general claims."
)

KEYWORDS = [
    "Proton exchange membrane fuel cell",
    "Fault classification",
    "Electrochemical impedance spectroscopy",
    "UniShape",
    "Prototype learning",
    "Kolmogorov-Arnold network",
    "Radial basis function",
]

ABBREVIATIONS = [
    ["Abbreviation", "Full term", "Meaning in this manuscript"],
    ["AdamW", "Adam with decoupled weight decay", "Optimizer used for model training"],
    ["CAPT", "Condition-Aware Prototype Transport", "Condition-driven dynamic prototype classification mechanism"],
    ["CE", "Cross-entropy", "Main supervised classification loss term"],
    ["CNN", "Convolutional neural network", "Baseline model family"],
    ["EIS", "Electrochemical impedance spectroscopy", "Impedance-related diagnostic information source"],
    ["GELU", "Gaussian error linear unit", "Activation function used in the neural modules"],
    ["KAN", "Kolmogorov-Arnold network", "Nonlinear residual branch used in Residual KAN-Fusion"],
    ["LN", "Layer normalization", "Normalization operation used in the condition encoder and fusion modules"],
    ["ML", "Machine learning", "Traditional baseline category"],
    ["MLP", "Multi-layer perceptron", "Feed-forward neural network module or baseline"],
    ["NPZ", "NumPy compressed archive", "Data file format used by the official experiment pipeline"],
    ["PEMFC", "Proton exchange membrane fuel cell", "Target fuel-cell system for fault classification"],
    ["RBF", "Radial basis function", "Condition mapper used to generate prototype offsets"],
    ["SNR", "Signal-to-noise ratio", "Noise level used in robustness experiments"],
    ["Macro-F1", "Macro-averaged F1 score", "Unweighted average F1 across classes"],
    ["Weighted-F1", "Support-weighted F1 score", "F1 score averaged by class support"],
]


SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Introduction",
        [
            "Proton exchange membrane fuel cells (PEMFCs) are attractive power sources for transportation, distributed generation, and portable energy systems because of their high conversion efficiency, rapid dynamic response, and low local emissions [1-3]. Their practical reliability, however, is constrained by tightly coupled electrochemical, thermal, water-management, and gas-transport processes. When the operating condition departs from a suitable range, the stack may enter states such as membrane dehydration, water flooding, reactant starvation, or other condition-dependent degradation patterns [1-3]. These faults can appear as voltage fluctuations, power loss, impedance changes, or shifts in operating statistics. Accurate fault classification is therefore a prerequisite for online monitoring, control intervention, and maintenance planning.",
            "A central difficulty in PEMFC fault diagnosis is that the same physical fault can present different observable signatures under different load, humidity, and stack-current conditions. Conversely, changes caused only by benign operating-condition drift may resemble early fault patterns. A model that ignores condition information may therefore confuse class-intrinsic fault structure with condition-induced distribution shift. This issue is especially important when the available dataset is limited and class supports are imbalanced, because the classifier may overfit to condition-specific patterns in a single split.",
            "Operational stack variables and electrochemical impedance-related features provide complementary diagnostic information. Stack voltage, current, and power describe the macroscopic temporal response of the system, while EIS or impedance-derived shape sequences encode information related to ohmic resistance, charge-transfer behavior, and mass-transport effects [4-6]. In the present project, the model input explicitly separates these sources into x_op, x_eis, and x_cond. This design avoids reducing all information to a single flat feature vector, while still allowing the model to learn cross-modal interactions.",
            "Recent time-series representation models, including shape-oriented backbone models such as UniShape, provide a useful basis for extracting transferable shape embeddings from sequential data [7]. Nevertheless, direct fine-tuning of a generic time-series classifier does not by itself address the condition drift problem. For PEMFC fault classification, a useful model should not only extract operation and impedance shapes, but also adapt its decision geometry according to the current operating condition.",
            "This paper proposes CAPT-UniShape, a Condition-Aware Prototype Transport UniShape framework. CAPT-UniShape uses two official UniShape backbones to encode the operation and EIS/impedance sequences, a condition encoder to embed x_cond, a Residual KAN-Fusion module to model nonlinear interactions among the three embeddings, and a condition-aware RBF prototype transport head to generate dynamic class prototypes. In the evaluated seed-44 split, this design obtains 100.00% clean-test accuracy under the fixed test protocol, but the analysis is deliberately conservative because the test set contains only 55 windows and the noise experiments reveal sensitivity of the normal class.",
            "The contributions of this paper are as follows. First, a three-input PEMFC fault-classification formulation is established using operation sequences, EIS/impedance shape sequences, and condition/statistical vectors. Second, CAPT-UniShape integrates official UniShape feature extraction with Residual KAN-Fusion and condition-aware RBF prototype transport. Third, a controlled ablation study is reported to separate the effects of the RBF head, Residual KAN-Fusion, prototype regularization, and input branches. Fourth, joint-input noise experiments are discussed explicitly, including failure modes under noisy conditions, rather than presenting the clean-test result as a broad robustness claim.",
        ],
    ),
    (
        "2. Related work",
        [
            "2.1. PEMFC fault diagnosis. PEMFC fault diagnosis has been studied using model-based observers, signal-processing features, traditional machine learning, and deep neural networks [1-3]. Voltage-based and current-based methods are easy to deploy, but their signals can be highly load-dependent. EIS-based methods are more directly connected to internal electrochemical processes and can distinguish changes in ohmic, charge-transfer, and mass-transport regions [4-6]. However, EIS measurements or impedance-derived features are often used as isolated features rather than as structured shape sequences, which may discard information in their ordered frequency or constructed-response pattern.",
            "2.2. Time-series shape representation. Deep learning models for time-series classification include convolutional networks, recurrent networks, Transformers, and more recent foundation-style time-series backbones [7]. Shape-oriented models are relevant for PEMFC diagnosis because faults may be expressed through local trends, response curvature, and multi-scale changes rather than only through scalar statistics. UniShape is used in this work as the official backbone for operation and EIS/impedance shape encoding [7]. The present study does not claim to improve the UniShape backbone itself; instead, it adapts the backbone to a multi-input PEMFC diagnosis setting.",
            "2.3. Multi-modal fusion under operating-condition drift. Multi-modal fault diagnosis commonly combines sensor streams, handcrafted features, and contextual variables through early fusion, late fusion, attention, or gating mechanisms [1-3]. For PEMFCs, condition variables are not merely auxiliary metadata. They define the operating envelope under which a fault signature is observed. CAPT-UniShape therefore treats x_cond as both a representation source in the fused feature vector and a driver of prototype transport in the classification head.",
            "2.4. Prototype learning, RBF mapping, and KAN components. Prototype-based classification represents each class by one or more reference vectors and classifies samples by similarity to those vectors [8]. Static prototypes can be interpretable, but they may be too rigid under condition-dependent class drift. RBF mappings provide smooth local interpolation from condition embeddings to prototype offsets [9,10]. KAN-style layers model nonlinear scalar transformations through basis expansions and can be used as compact residual nonlinear modules [11]. In CAPT-UniShape, these ideas are used pragmatically: the RBF mapper produces condition-dependent prototype offsets, while the KAN branch is placed behind a bottleneck and added as a residual to a stable MLP fusion branch.",
        ],
    ),
    (
        "3. Methodology",
        [
            "3.1. Problem definition. Let D = {(x_i^op, x_i^eis, x_i^cond, y_i)}_{i=1}^N denote a labeled PEMFC dataset with K fault classes. In the evaluated dataset, x_i^op in R^{3 x 64} contains stack total voltage, stack total current, and stack power; x_i^eis in R^{4 x 128} is a constructed EIS/impedance statistical shape sequence; x_i^cond in R^{12} is a condition/statistical vector; and y_i in {0, ..., K-1}. Class 0 is Normal, class 1 is reported as Drying/membrane dehydration according to the current paper materials, and class 2 is reported as Flooding/over-wet according to the author's confirmation. The objective is to learn a classifier f_theta that maps the three-input tuple to class logits and predicted label y_hat_i = argmax_k f_theta(x_i)_k.",
            "3.2. Constructed EIS/impedance shape sequence. The present x_eis input is not a raw full-frequency EIS spectrum. It is constructed from nine impedance/EIS statistical variables: total impedance, mean impedance, maximum impedance, second-highest impedance, minimum impedance, second-lowest impedance, standard deviation, EIS resistance real part, and EIS resistance imaginary part. Let s_i in R^9 denote this ordered statistic vector. The builder interpolates s_i to a length-128 curve c_i, computes its first-difference shape g_i, computes a centered cumulative shape a_i, and appends a normalized coordinate q in [0,1]. Thus",
            "(1)    x_i^eis = [c_i; g_i; a_i; q] in R^{4 x 128}.",
            "This representation preserves an ordered impedance-statistical shape interface for UniShape while avoiding the unsupported claim that full raw EIS frequency spectra are available in the current dataset.",
            "3.3. UniShape branch encoders. The operation and EIS/impedance branches use official UniShape backbone wrappers. For multi-channel inputs, each channel is encoded by the shared official UniShape encoder and then aggregated across channels. Denote the two feature extractors as B_op and B_eis. The embeddings are",
            "(2)    z_op = B_op(x_op),      z_eis = B_eis(x_eis),      z_op, z_eis in R^d.",
            "This design keeps the sequential structure of both operation and impedance-shape inputs, instead of flattening them before representation learning.",
            "3.4. Condition encoder. The condition/statistical vector is encoded by an MLP condition encoder phi_cond. It maps the 12-dimensional vector into the same representation dimension d as the UniShape embeddings:",
            "(3)    z_cond = phi_cond(x_cond) = W_2 GELU(LN(W_1 x_cond + b_1)) + b_2.",
            "The encoded condition token is used in two places. It is concatenated with the two branch embeddings for feature fusion, and it drives the prototype-offset mapper in the classification head.",
            "3.5. Residual KAN-Fusion. The three embeddings are concatenated as u = [z_op; z_eis; z_cond] in R^{3d}. CAPT-UniShape fuses them using Residual KAN-Fusion:",
            "(4)    h = MLP(u) + lambda_KAN * W_K KAN(Bottleneck(u)).",
            "The MLP branch provides a stable primary fusion path. The KAN branch first compresses u through a bottleneck before applying a KAN-style nonlinear layer, then projects the result back to the fused dimension. This keeps the nonlinear branch controlled and prevents the KAN layer from directly operating on the full high-dimensional concatenated input. The scalar lambda_KAN is learnable in the implemented model.",
            "The KAN-style layer used in this project expands normalized scalar inputs on learnable Gaussian basis functions and linearly mixes the expanded basis responses. Its regularization penalizes basis-weight scale and center spacing, yielding a compact nonlinear residual rather than a standalone classifier.",
            "3.6. Condition-aware RBF prototype transport head. Let P^0 = {p_k^0}_{k=1}^K be the learnable static class prototypes, with p_k^0 in R^d. The condition-aware head computes a dynamic prototype for each class and sample:",
            "(5)    p_{i,k} = p_k^0 + Delta p_{i,k}(z_i^cond).",
            "The offset Delta p_{i,k} is generated by an RBF condition mapper. Given RBF centers c_j and widths sigma_j, the j-th RBF response is",
            "(6)    r_{i,j} = exp(-||z_i^cond - c_j||_2^2 / (2 sigma_j^2)).",
            "The vector r_i is linearly mapped and reshaped into K prototype offsets. The class logit is then computed by cosine similarity with temperature tau:",
            "(7)    logit_{i,k} = cos(h_i, p_{i,k}) / tau.",
            "The resulting classifier can be interpreted as a condition-adapted prototype classifier. It does not require each fault class to collapse to one fixed point under all operating conditions; instead, it permits a controlled condition-driven displacement around a static class anchor.",
            "3.7. Training objective. The training loss combines class-weighted cross-entropy with prototype-transport, prototype-separation, and KAN regularization terms:",
            "(8)    L = L_CE + alpha_transport L_transport + alpha_sep L_sep + alpha_KAN L_KAN.",
            "In the implemented configuration, alpha_transport = 0.001, alpha_sep = 0.001, and alpha_KAN = 0.0001. The cross-entropy term uses sqrt-balanced class weighting in the reported experiments. The transport penalty constrains prototype offsets by their squared magnitude:",
            "(9)    L_transport = mean_i,k ||Delta p_{i,k}||_2^2.",
            "The separation term penalizes excessive cosine similarity among static prototypes using a margin m:",
            "(10)    L_sep = mean_{a != b} max(0, cos(p_a^0, p_b^0) - m).",
            "The KAN regularizer L_KAN is the scale and smoothness penalty returned by the KAN branch. During prediction, only the logits in Eq. (7) are used.",
        ],
    ),
    (
        "4. Experimental setup",
        [
            "4.1. Data source, labels, and recorded variables. The evaluated project data are stored in an Excel source file and converted into an NPZ dataset for the official model pipeline. The raw table contains 11,137 rows and 230 columns. Available columns include test time, numeric label, 216 single-cell voltage columns, stack total voltage, stack total current, stack power, and nine impedance/EIS statistical variables. The original row-level label counts are class 0: 1,752, class 1: 2,036, and class 2: 7,349. In this manuscript, class 0 is Normal, class 1 is reported as Drying/membrane dehydration, and class 2 is reported as Flooding/over-wet according to the author's confirmation that code 2 corresponds to over-wet data.",
            "After windowing, the evaluated 8:2 dataset contains 433 windows with labels distributed as class 0: 74, class 1: 68, and class 2: 291. The 8:2 fixed test split contains 311 training windows, 67 validation windows, and 55 test windows. The test label counts are class 0: 8, class 1: 11, and class 2: 36.",
            "The model inputs are x_op in R^{N x 3 x 64}, x_eis in R^{N x 4 x 128}, and x_cond in R^{N x 12}. The operation channels are stack total voltage, stack total current, and stack power. The condition vector consists of nine impedance/EIS statistical features plus three stack-level variables. The EIS/impedance sequence is a four-channel constructed statistical shape sequence with length 128, not a raw full-frequency EIS spectrum. If full raw EIS spectra are used in a future revision, the manuscript should state that explicitly and update this description.",
            "The available raw-table ranges are summarized in Table 1. These values support a basic description of the recorded variables, but they do not replace a full experimental-platform description. Stack specification, test bench model, EIS/impedance instrument, sensor models, sampling frequency, gas supply, pressure, temperature, humidity, and flow-rate settings should be added by the authors before submission.",
            "The fixed test protocol uses seed = 44, window_size = 64, stride_train = 16, stride_eval = 32, eis_seq_len = 128, split_mode = segment, segment_block_seconds = 300, ratio = 8:2, and val_size = 0.25. Validation is used for checkpoint selection, and the final reported model uses best.ckpt. Train+validation refitting is not used in the reported main result.",
            "4.2. Training details. The official CAPT-UniShape-RBF-KANFusion model is trained with AdamW, learning rate 0.0001, weight decay 0.0001, batch size 32, maximum 80 epochs, early-stopping patience 10, and sqrt-balanced class weighting. The shared representation dimension is d = 128. The RBF head uses 16 RBF centers and an initial temperature of 0.07. The Residual KAN-Fusion bottleneck dimension is 32, the number of KAN basis functions is 8, and lambda_KAN is initialized to 0.1.",
            "4.3. Compared models and ablations. The available clean 8:2 project artifacts include traditional machine-learning baselines (logistic regression and random forest), neural baselines (MLP and 1D CNN), Transformer-style baselines (Transformer and iTransformer), and the proposed Official-CAPT-UniShape-RBF-KANFusion. The ablation study includes no_rbf, no_kan_fusion, static_prototype, no_transport_reg, no_separation_reg, no_eis_input, no_condition_input, stack_only, and eis_cond_only. These ablations are designed to isolate the roles of the prototype head, residual KAN fusion, regularizers, and input modalities.",
            "4.4. Evaluation metrics. The primary metrics are test accuracy, Macro-F1, and Weighted-F1. Macro-F1 is emphasized because the test set is class-imbalanced and the normal class has only eight test windows. Per-class precision, recall, and F1 are reported for class 0 because false alarms and missed normal-state recognition are important for diagnostic reliability. Parameter count is reported as a model-size indicator.",
        ],
    ),
    (
        "5. Results and discussion",
        [
            "5.1. Main clean-test performance. Under the evaluated seed-44 fixed 8:2 test protocol, Official-CAPT-UniShape-RBF-KANFusion achieved 100.00% test accuracy, 100.00% Macro-F1, and 100.00% Weighted-F1 on 55 test windows. The class-0 normal category also achieved 100.00% precision, recall, and F1 in this clean-test split. The model contains 6,489,373 parameters.",
            "This result should be interpreted as strong performance under the fixed evaluated split, not as evidence of broad generalization. The test set is small, class 0 has only eight test windows, and the current result is based on a single seed. Nevertheless, the clean-test confusion matrix indicates that the model can separate the three encoded classes under the current segment-based protocol.",
            "5.2. Contextual comparison with available baselines. The available 8:2 baseline artifacts show that several simpler models also perform strongly on the clean split. Logistic regression and random forest both reach 94.55% accuracy and 88.31% Macro-F1, while the Transformer baseline reaches 94.55% accuracy and 93.29% Macro-F1 in the aligned clean summary. This suggests that the dataset contains highly informative EIS/condition-related patterns. Therefore, the contribution of CAPT-UniShape should be framed in terms of multi-input modeling, dynamic prototype formulation, and ablation behavior rather than only as a larger absolute accuracy gain.",
            "5.3. Ablation analysis. The ablation results are more informative than the clean-test headline metric. Removing the RBF dynamic prototype head decreases test accuracy from 100.00% to 98.18% and Macro-F1 from 100.00% to 96.33%; class-0 recall drops to 87.50%. Removing Residual KAN-Fusion gives a larger decline, with accuracy 96.36%, Macro-F1 92.46%, and class-0 recall 75.00%. These results indicate that the RBF prototype head and the KAN residual branch are useful for maintaining normal-class sensitivity in the evaluated split.",
            "However, not every subcomponent yields a visible marginal gain on this split. The static_prototype, no_transport_reg, and no_separation_reg variants all reach 100.00% accuracy and 100.00% Macro-F1. Therefore, the current data do not justify a strong claim that condition transport regularization or prototype separation regularization is necessary for clean-test performance. Their role should be treated as a regularization design that may be useful in broader settings, pending further validation.",
            "5.4. Input-branch sensitivity. The input ablations show that the EIS/impedance branch is the most important source of discriminative information in the current dataset. When the EIS branch is removed, accuracy falls to 87.27%, Macro-F1 to 82.99%, and class-0 F1 to 58.82%. When only the stack operation branch is retained, accuracy falls further to 83.64% and Macro-F1 to 79.95%. By contrast, the eis_cond_only variant achieves 100.00% accuracy and 100.00% Macro-F1. This does not mean that operation signals are useless in general; rather, it shows that in the current evaluated split, the EIS/impedance and condition/statistical channels are sufficient for clean classification, while the operation branch adds limited marginal value.",
            "5.5. Noise robustness. The clean model is sensitive to joint perturbations applied to x_op, x_eis, and x_cond at test time. At 30 dB and 20 dB SNR, accuracy remains at 85.45% and 87.27%, respectively, but Macro-F1 decreases to 62.26% and 68.32%. The class-0 normal category is particularly sensitive: at 30 dB, class-0 precision, recall, and F1 are all 0.00%; at 20 dB, class-0 recall is 12.50% and F1 is 22.22%. Stronger perturbations cause larger degradation, including 83.64% accuracy at 10 dB, 60.00% at 5 dB, and 25.45% at 0 dB.",
            "These results should be discussed carefully. CAPT-UniShape performs very well under the clean fixed protocol, but the present experiments do not support a claim that the method is robust under all noise strengths or consistently superior to all baselines under noise. The observed failure mode is aligned with the small support of class 0 and the reliance on EIS/condition patterns. Future work should evaluate noise-aware training, modality-specific denoising, calibration, and repeated multi-seed validation.",
        ],
    ),
    (
        "6. Limitations",
        [
            "The current study has several limitations. First, the main reported result is based on a single seed, seed 44. Multi-seed reporting is required before making stable claims about average performance or variance. Second, the fixed test set contains only 55 windows, and the normal class contains only eight test windows. A 100.00% clean-test score under this protocol may therefore be sensitive to the split. Third, the evaluated dataset contains three encoded classes; external datasets, additional PEMFC fault modes, and cross-condition transfer settings are not yet demonstrated in this manuscript.",
            "Fourth, the noise experiments show that the normal class is sensitive to joint perturbations of operation, EIS/impedance, and condition inputs. Strong-noise robustness remains insufficient. Fifth, several prototype-related ablations, including static_prototype, no_transport_reg, and no_separation_reg, also reach 100.00% in the clean split. This means that their marginal contribution cannot be overstated using the current evidence. Sixth, the EIS/impedance sequence used in this project is constructed from impedance/EIS statistics; if future experiments use raw full-frequency EIS spectra, the data description and claims should be updated accordingly.",
        ],
    ),
    (
        "7. Conclusion",
        [
            "This paper presented CAPT-UniShape, a condition-aware prototype transport framework for PEMFC fault classification. The model combines official UniShape operation and EIS/impedance shape encoders, an MLP condition encoder, Residual KAN-Fusion, and a condition-aware RBF prototype transport head. Under the evaluated seed-44 fixed 8:2 test protocol, the model achieved 100.00% accuracy, 100.00% Macro-F1, and 100.00% Weighted-F1 on 55 test windows. Ablation results show that the EIS/impedance branch is critical in the present dataset, and that RBF prototypes and Residual KAN-Fusion improve normal-class recall relative to their removal.",
            "The findings should be interpreted with appropriate caution. The current evidence supports CAPT-UniShape as a promising multi-input architecture for the evaluated PEMFC dataset, but not as a broadly robust diagnostic solution. Broader validation should include multi-seed experiments, larger test sets, external operating conditions, raw EIS spectra where available, and noise-aware training strategies.",
        ],
    ),
    (
        "Declaration of competing interest",
        [
            "The authors should declare whether they have any known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.",
        ],
    ),
    (
        "Data availability",
        [
            "The data availability statement should be completed by the authors. The present draft is based on project files and local experimental artifacts in the CAPT-UniShape workspace.",
        ],
    ),
    (
        "References",
        [
            "[1] Araya, S. S., Zhou, F., Sahlin, S. L., Thomas, S., Jeppesen, C., and Kaer, S. K. Fault characterization of a proton exchange membrane fuel cell stack. Energies, 2019, 12(1), 152. https://doi.org/10.3390/en12010152.",
            "[2] Benmouna, A., Becherif, M., Depernet, D., Gustin, F., Ramadan, H. S., and Fukuhara, S. Fault diagnosis methods for Proton Exchange Membrane Fuel Cell system. International Journal of Hydrogen Energy, 2017, 42(2), 1534-1543. https://doi.org/10.1016/j.ijhydene.2016.07.181.",
            "[3] Wang, Y., et al. Water management fault diagnosis for proton-exchange membrane fuel cells based on deep learning methods. International Journal of Hydrogen Energy, 2023, 48(72), 28163-28173. https://doi.org/10.1016/j.ijhydene.2023.03.097.",
            "[4] Wasterlain, S., Candusso, D., Harel, F., Hissel, D., and Francois, X. Characterisation of proton exchange membrane fuel cell failures via electrochemical impedance spectroscopy. Journal of Power Sources, 2006, 161(1), 264-274. https://doi.org/10.1016/j.jpowsour.2006.03.067.",
            "[5] Ibrahim, M., et al. Rapid fault diagnosis of PEM fuel cells through optimal electrochemical impedance spectroscopy tests. Energies, 2020, 13(14), 3643. https://doi.org/10.3390/en13143643.",
            "[6] Online fault detection and isolation of PEMFC based on EIS and data-driven methods: Feasibility study and prospects. Journal of Power Sources, 2025. DOI/source page: https://www.sciencedirect.com/science/article/pii/S0378775325007517.",
            "[7] Liu, Z., Wang, Y., Li, B., Zheng, J., Eldele, E., Wu, M., and Ma, Q. A unified shape-aware foundation model for time series classification. arXiv:2601.06429, 2026. https://arxiv.org/abs/2601.06429.",
            "[8] Snell, J., Swersky, K., and Zemel, R. Prototypical networks for few-shot learning. NeurIPS, 2017. https://papers.neurips.cc/paper/6996-prototypical-networks-for-few-shot-learning.",
            "[9] Broomhead, D. S., and Lowe, D. Multivariable functional interpolation and adaptive networks. Complex Systems, 1988, 2, 321-355. https://www.complex-systems.com/abstracts/v02_i03_a05/.",
            "[10] Powell, M. J. D. The theory of radial basis function approximation in 1990. In Advances in Numerical Analysis: Wavelets, Subdivision Algorithms, and Radial Basis Functions, 1992, 105-210. https://doi.org/10.1093/oso/9780198534396.003.0003.",
            "[11] Liu, Z., Wang, Y., Vaidya, S., Ruehle, F., Halverson, J., Soljacic, M., Hou, T. Y., and Tegmark, M. KAN: Kolmogorov-Arnold Networks. arXiv:2404.19756, 2024. https://arxiv.org/abs/2404.19756.",
        ],
    ),
    (
        "Information that should be supplied by the authors",
        [
            "1. Confirm whether encoded class 1 should be described as Drying / membrane dehydration in all paper tables, figure labels, and configuration files. Class 2 has been set to Flooding / over-wet according to the current author instruction.",
            "2. Provide the experimental platform description, PEMFC stack specifications, sensor details, and operating-condition ranges.",
            "3. Keep the current wording that x_eis is a constructed EIS/impedance statistical shape sequence unless raw full-frequency EIS spectra are added later.",
            "4. Decide whether to include additional multi-seed results, external validation, or cross-condition validation before journal submission.",
            "5. Update all generated result figures so their class labels match the final label convention.",
            "6. Provide author names, affiliations, acknowledgements, funding information, conflict-of-interest statement, and data/code availability statement.",
        ],
    ),
]

def _strip_subsection_prefix(text: str) -> str:
    """Remove thesis-like subsection numbering while keeping equation labels."""
    import re

    return re.sub(r"^\d+\.\d+\.\s*", "", text)


MATERIALS_AND_METHODS_REWRITTEN = [
    "Data representation and diagnostic setting. The diagnostic task considered in this study is three-class PEMFC fault classification under a fixed segment-based evaluation protocol. The available data contain stack-level operation variables, impedance/EIS-derived statistics, and condition-related statistical variables. These sources describe different aspects of the same operating state: stack voltage, current, and power capture the macroscopic dynamic response of the system; impedance-related variables reflect electrochemical and transport behavior; and condition/statistical variables provide the operating context under which the observed response occurs. CAPT-UniShape is designed around this separation instead of collapsing all variables into a single flat feature vector.",
    "Let D = {(x_i^op, x_i^eis, x_i^cond, y_i)}_{i=1}^N denote the labeled dataset with K classes. For each window, x_i^op in R^{3 x 64} contains stack total voltage, stack total current, and stack power; x_i^eis in R^{4 x 128} is a constructed EIS/impedance statistical shape sequence; x_i^cond in R^{12} is a condition/statistical vector; and y_i in {0, ..., K - 1} is the class label. In the current manuscript, class 0 denotes Normal, class 1 is reported as Drying/membrane dehydration according to the available paper materials, and class 2 denotes Flooding/over-wet according to the author's confirmation. The model learns a mapping f_theta from the three inputs to class logits, and the predicted class is obtained as y_hat_i = argmax_k f_theta(x_i)_k.",
    "Construction of the EIS/impedance shape input. The input x_eis used in this work should be interpreted as a constructed impedance-statistical shape sequence, not as a raw full-frequency EIS spectrum. It is generated from nine ordered impedance/EIS statistical variables: total impedance, mean impedance, maximum impedance, second-highest impedance, minimum impedance, second-lowest impedance, standard deviation, EIS resistance real part, and EIS resistance imaginary part. Let s_i in R^9 denote this ordered vector. The builder interpolates s_i into a length-128 curve c_i, computes its first-difference sequence g_i, computes a centered cumulative shape a_i, and appends a normalized coordinate q in [0, 1]. The resulting four-channel sequence is",
    "(1)    x_i^eis = [c_i; g_i; a_i; q] in R^{4 x 128}.",
    "where x_i^eis denotes the constructed EIS/impedance shape sequence of the i-th window; c_i, g_i, and a_i denote the interpolated statistic curve, first-difference sequence, and centered cumulative shape, respectively; q is the normalized coordinate channel; and R^{4 x 128} specifies the four-channel sequence length.",
    "This construction provides a shape-like interface for the UniShape encoder while keeping the data description consistent with the available project files. If raw frequency-resolved EIS spectra are introduced in a later revision, the data description and corresponding claims should be updated accordingly.",
    "CAPT-UniShape architecture. The model follows a three-branch representation design. The operation sequence and the constructed EIS/impedance shape sequence are encoded by two official UniShape backbone wrappers. For multi-channel inputs, each channel is passed through the shared official UniShape encoder and the channel-level representations are aggregated. Denoting the two branch encoders by B_op and B_eis, the sequential embeddings are",
    "(2)    z_i^op = B_op(x_i^op),      z_i^eis = B_eis(x_i^eis),      z_i^op, z_i^eis in R^d.",
    "where B_op and B_eis are the operation and EIS/impedance UniShape branch encoders; x_i^op and x_i^eis are their corresponding inputs; z_i^op and z_i^eis are the learned branch embeddings; and d is the embedding dimension.",
    "The condition/statistical vector is encoded by an MLP condition encoder phi_cond, which maps the 12-dimensional input into the same representation dimension d:",
    "(3)    z_i^cond = phi_cond(x_i^cond) = W_2 GELU(LN(W_1 x_i^cond + b_1)) + b_2.",
    "where x_i^cond is the 12-dimensional condition/statistical vector; phi_cond denotes the condition encoder; W_1 and W_2 are linear projection matrices; b_1 and b_2 are bias terms; LN denotes layer normalization; GELU is the Gaussian error linear unit; and z_i^cond is the encoded condition token.",
    "The condition embedding is used both as a feature source in the fusion module and as the driver of condition-dependent prototype displacement in the classification head. This design reflects the assumption that operating conditions can influence not only the feature distribution but also the appropriate class decision geometry.",
    "Residual KAN-Fusion. The three embeddings are concatenated as u_i = [z_i^op; z_i^eis; z_i^cond] in R^{3d}. CAPT-UniShape fuses them using a residual structure composed of a stable MLP branch and a bottleneck KAN residual branch:",
    "(4)    h_i = MLP(u_i) + lambda_KAN W_K KAN(Bottleneck(u_i)).",
    "where u_i is the concatenated multi-source embedding; h_i is the fused representation; MLP is the main fusion branch; Bottleneck compresses u_i before the KAN branch; KAN denotes the Kolmogorov-Arnold network residual transformation; W_K is the projection matrix after the KAN branch; and lambda_KAN is the learnable residual scaling coefficient.",
    "The MLP branch forms the main fusion path. The KAN branch first compresses the concatenated representation through a bottleneck, applies a KAN-style nonlinear transformation, and then projects the output back to the fused representation dimension. In the implemented model, lambda_KAN is learnable and is initialized to 0.1. The KAN-style layer expands normalized scalar inputs on learnable Gaussian basis functions and linearly mixes the basis responses. Its regularization penalizes basis-weight scale and center spacing, so the branch is used as a controlled nonlinear residual rather than as an independent classifier.",
    "Condition-aware RBF prototype transport head. The classifier is formulated as a condition-adapted prototype classifier. Let P^0 = {p_k^0}_{k=1}^K denote learnable static class prototypes, where p_k^0 in R^d. For each sample and class, the head computes a dynamic prototype by adding a condition-dependent offset:",
    "(5)    p_{i,k} = p_k^0 + Delta p_{i,k}(z_i^cond).",
    "where p_k^0 is the static prototype of class k; Delta p_{i,k}(z_i^cond) is the condition-dependent prototype offset generated for sample i and class k; and p_{i,k} is the resulting dynamic prototype.",
    "The offset is produced by an RBF condition mapper. Given RBF centers c_j and widths sigma_j, the j-th response is",
    "(6)    r_{i,j} = exp(-||z_i^cond - c_j||_2^2 / (2 sigma_j^2)).",
    "where r_{i,j} is the j-th RBF response for sample i; z_i^cond is the condition token; c_j is the j-th RBF center; sigma_j is its width; and ||.||_2 denotes the Euclidean norm.",
    "The RBF response vector is linearly mapped and reshaped into class-wise prototype offsets. The class logit is then computed by temperature-scaled cosine similarity:",
    "(7)    logit_{i,k} = cos(h_i, p_{i,k}) / tau.",
    "where logit_{i,k} is the classification logit for sample i and class k; cos(.) denotes cosine similarity; h_i is the fused representation; p_{i,k} is the dynamic class prototype; and tau is the temperature parameter.",
    "Compared with a purely static prototype head, this formulation allows each class anchor to move in a restricted, condition-dependent manner. The claim made here is deliberately limited: the head is intended to model condition-related decision shifts in the evaluated dataset, and its necessity should be judged by the ablation results rather than assumed a priori.",
    "Training objective. The model is trained with a class-weighted cross-entropy term and three auxiliary regularizers:",
    "(8)    L = L_CE + alpha_transport L_transport + alpha_sep L_sep + alpha_KAN L_KAN.",
    "where L is the total training loss; L_CE is the class-weighted cross-entropy loss; L_transport is the prototype-transport regularization term; L_sep is the prototype-separation regularization term; L_KAN is the KAN regularization term; and alpha_transport, alpha_sep, and alpha_KAN are the corresponding loss weights.",
    "In the reported configuration, alpha_transport = 0.001, alpha_sep = 0.001, and alpha_KAN = 0.0001. The cross-entropy term uses sqrt-balanced class weighting. The transport penalty constrains the magnitude of condition-driven prototype offsets,",
    "(9)    L_transport = mean_{i,k} ||Delta p_{i,k}||_2^2,",
    "where mean_{i,k} denotes averaging over samples and classes, and Delta p_{i,k} is the condition-dependent prototype offset.",
    "and the separation term penalizes excessive similarity among static prototypes using a margin m:",
    "(10)    L_sep = mean_{a != b} max(0, cos(p_a^0, p_b^0) - m).",
    "where mean_{a != b} denotes averaging over all unequal class-prototype pairs; p_a^0 and p_b^0 are static prototypes of classes a and b; cos(.) is cosine similarity; and m is the separation margin.",
    "The KAN regularizer L_KAN is the scale and smoothness penalty returned by the KAN branch. During inference, only the logits in Eq. (7) are used.",
    "Dataset construction and split protocol. The evaluated project data are stored in an Excel source file and converted into an NPZ dataset for the official model pipeline. The raw table contains 11,137 rows and 230 columns, including test time, numeric label, 216 single-cell voltage columns, stack total voltage, stack total current, stack power, and nine impedance/EIS statistical variables. The row-level label counts are 1,752 for class 0, 2,036 for class 1, and 7,349 for class 2. After windowing, the evaluated 8:2 dataset contains 433 windows, with 74 windows from class 0, 68 from class 1, and 291 from class 2.",
    "The fixed test protocol uses seed = 44, window_size = 64, stride_train = 16, stride_eval = 32, eis_seq_len = 128, split_mode = segment, segment_block_seconds = 300, ratio = 8:2, and val_size = 0.25. The resulting split contains 311 training windows, 67 validation windows, and 55 test windows. The test set contains 8 Normal windows, 11 class-1 windows, and 36 Flooding/over-wet windows. Validation is used for checkpoint selection, the reported model uses best.ckpt, and train+validation refitting is not used.",
    "Training configuration, baselines, and metrics. Official-CAPT-UniShape-RBF-KANFusion is trained with AdamW, learning rate 0.0001, weight decay 0.0001, batch size 32, a maximum of 80 epochs, early-stopping patience 10, and sqrt-balanced class weighting. The shared representation dimension is d = 128. The RBF head uses 16 RBF centers and an initial temperature of 0.07. The Residual KAN-Fusion bottleneck dimension is 32 and the KAN branch uses 8 basis functions.",
    "The available clean 8:2 comparison results include logistic regression, random forest, MLP, 1D CNN, Transformer, iTransformer, and the proposed model. The ablation variants include no_rbf, no_kan_fusion, static_prototype, no_transport_reg, no_separation_reg, no_eis_input, no_condition_input, stack_only, and eis_cond_only. Accuracy, Macro-F1, and Weighted-F1 are used as the main metrics. Macro-F1 is emphasized because the evaluated test set is imbalanced and the Normal class contains only eight test windows. Class-0 precision, recall, and F1 are also reported to make the normal-state behavior explicit. The raw-variable ranges, clean fixed-protocol model summary, and available baseline results are provided in Tables 1-3 to support reproducibility of the experimental setting; the interpretation of these results is discussed in Section 3.",
]


RESULTS_AND_DISCUSSION_REWRITTEN = [
    "Scope of evaluation. The results in this section should be interpreted within the fixed experimental protocol described in Section 2. The main model is evaluated with seed = 44, an 8:2 segment-based split, validation-based selection of best.ckpt, and no train+validation refitting. The final test set contains 55 windows, including only 8 Normal windows, 11 class-1 windows, and 36 Flooding/over-wet windows. This setting is useful for comparing model variants under a controlled project protocol, but it is not sufficient by itself to support broad claims about cross-seed, cross-dataset, or deployment-level generalization.",
    "Clean-test performance under the fixed protocol. Under the evaluated seed-44 fixed 8:2 test protocol, Official-CAPT-UniShape-RBF-KANFusion achieved 100.00% accuracy, 100.00% Macro-F1, and 100.00% Weighted-F1 on the 55-window test set (Table 2). The class-0 Normal category also reached 100.00% precision, recall, and F1 in this clean-test split. The corresponding confusion matrix in Fig. 3 shows no misclassified test windows among the three encoded classes. These results indicate that the proposed three-input representation and condition-aware classification design can separate the available test windows in this evaluated split.",
    "The 100.00% result should not be read as evidence that the model is error-free in general. The score is obtained from a single seed and a small fixed test set, and the Normal class has limited support. A single misclassification in class 0 would substantially change the class-level recall and F1. Therefore, the clean-test result is best framed as strong split-specific evidence rather than as a general performance guarantee.",
    "Comparison with available clean baselines. The available aligned 8:2 baseline results in Table 3 show that the dataset is informative even for simpler methods. Logistic regression and random forest both achieved 94.55% accuracy and 88.31% Macro-F1, while the Transformer baseline achieved 94.55% accuracy and 93.29% Macro-F1. The proposed model obtains the best clean-test scores among the available artifacts, but the absolute gap should be discussed with caution because all results are from the same fixed seed-44 protocol. The more informative evidence for the proposed design is therefore provided by the ablation behavior and the modality-sensitivity analysis, rather than by the headline accuracy alone.",
    "Effect of the prototype head and Residual KAN-Fusion. The ablation results in Table 4 support a moderate contribution from the condition-aware RBF prototype head and the Residual KAN-Fusion module in the evaluated split. Removing the RBF dynamic prototype head reduced accuracy from 100.00% to 98.18% and Macro-F1 from 100.00% to 96.33%; class-0 recall decreased from 100.00% to 87.50%. Removing Residual KAN-Fusion caused a larger degradation, with 96.36% accuracy, 92.46% Macro-F1, and 75.00% class-0 recall. These results suggest that the dynamic prototype formulation and nonlinear residual fusion are useful for maintaining Normal-class sensitivity under the current clean protocol.",
    "At the same time, the ablation evidence does not justify an overstatement of every subcomponent. The static_prototype, no_transport_reg, and no_separation_reg variants all achieved 100.00% accuracy and 100.00% Macro-F1 in the clean seed-44 split. Thus, the current evidence does not prove that condition-dependent prototype transport, transport regularization, or separation regularization is individually necessary for clean-test success. These components should be presented as part of a plausible condition-aware design whose marginal value requires further validation, especially across additional seeds and external test conditions.",
    "Input-branch sensitivity. The input ablations show that the EIS/impedance and condition/statistical information carries most of the discriminative signal in the current dataset. Removing the EIS/impedance branch reduced accuracy to 87.27%, Macro-F1 to 82.99%, and class-0 F1 to 58.82%. The stack_only variant achieved 83.64% accuracy and 79.95% Macro-F1, indicating that stack voltage, current, and power alone are less sufficient under this split. By contrast, the eis_cond_only variant reached 100.00% accuracy and 100.00% Macro-F1. This pattern should not be interpreted as evidence that operation sequences are generally unnecessary for PEMFC diagnosis. It indicates that, in the evaluated dataset and split, the constructed EIS/impedance shape sequence and condition/statistical vector are highly informative, while the operation branch adds limited visible marginal gain in the clean setting.",
    "Noise robustness and class-level failure mode. The joint-input noise experiments expose a different behavior from the clean protocol. When noise is added simultaneously to x_op, x_eis, and x_cond, the model no longer maintains balanced class performance. At 30 dB SNR, accuracy remained 85.45%, but Macro-F1 decreased to 62.26% and class-0 precision, recall, and F1 were all 0.00%. At 20 dB SNR, accuracy was 87.27%, but Macro-F1 was 68.32%, with class-0 recall of 12.50% and class-0 F1 of 22.22%. Stronger perturbations caused further degradation, including 60.00% accuracy at 5 dB and 25.45% accuracy at 0 dB (Table 5 and Fig. 5).",
    "These results indicate that CAPT-UniShape performs strongly under the clean fixed protocol but remains sensitive to joint input perturbations, particularly for the Normal class. The discrepancy between accuracy and Macro-F1 under noise is important because the test set is imbalanced: acceptable-looking accuracy can coexist with poor recognition of the minority Normal class. The noise experiments therefore should be reported as a limitation and diagnostic failure mode, not as evidence of broad robustness. They also suggest that future work should consider noise-aware training, modality-specific perturbation analysis, calibration, and repeated evaluation across multiple random seeds.",
    "Overall interpretation. Taken together, the experiments support a restrained conclusion. CAPT-UniShape is a useful multi-input architecture for the evaluated PEMFC dataset because it combines operation sequences, constructed EIS/impedance shape sequences, and condition/statistical information, and because its full configuration performs best in the clean seed-44 split. The ablations further indicate that EIS/impedance information is essential in the present dataset and that RBF prototypes and Residual KAN-Fusion help preserve Normal-class recall relative to their removal. However, the evidence is limited by the single seed, the 55-window test set, the small Normal-class support, and the observed noise sensitivity. Multi-seed testing, larger external test sets, and statistical comparison are needed before making stronger claims about general diagnostic robustness.",
]


CONCLUSIONS_REWRITTEN = [
    "This study presented CAPT-UniShape, a condition-aware prototype transport framework for PEMFC fault classification. The work addresses a practical diagnostic setting in which stack operation sequences, constructed EIS/impedance shape sequences, and condition/statistical variables provide complementary but condition-dependent information. CAPT-UniShape combines two official UniShape branch encoders, an MLP condition encoder, Residual KAN-Fusion, and a condition-aware RBF prototype transport head to integrate these inputs and adapt the prototype-based decision geometry to the encoded operating condition.",
    "Under the evaluated seed-44 fixed 8:2 test protocol, Official-CAPT-UniShape-RBF-KANFusion achieved 100.00% accuracy, 100.00% Macro-F1, and 100.00% Weighted-F1 on 55 test windows. This result shows that the proposed configuration can separate the three encoded classes in the clean fixed split. The ablation results further indicate that the constructed EIS/impedance branch is the most important input source in the current dataset, and that removing either the RBF dynamic prototype head or Residual KAN-Fusion reduces Normal-class recall. These findings support the usefulness of multi-input representation learning and condition-aware classification in the evaluated setting.",
    "The conclusions should remain limited to the available evidence. The main result is based on a single seed, a fixed 8:2 segment-based split, and a small test set with only eight Normal windows. In addition, static_prototype, no_transport_reg, and no_separation_reg also reached 100.00% accuracy and 100.00% Macro-F1 in the clean split, so the marginal contribution of these prototype-related regularization components cannot be claimed as decisive from the present experiments. The joint-input noise tests show a clear degradation of Macro-F1 and a pronounced sensitivity of the Normal class, indicating that the model should not be described as broadly robust to strong noise.",
    "Future work should therefore prioritize multi-seed evaluation, larger and external test sets, cross-condition validation, statistical significance testing, and noise-aware training or modality-specific denoising. If raw full-frequency EIS spectra become available, the current constructed EIS/impedance statistical shape input should also be compared with spectrum-based representations. Overall, the present results support CAPT-UniShape as a carefully scoped multi-source diagnostic model for the evaluated PEMFC dataset, rather than as a general-purpose PEMFC fault diagnosis solution.",
]


JOURNAL_SECTIONS: list[tuple[str, list[str]]] = [
    (
        "1. Introduction",
        SECTIONS[0][1] + [_strip_subsection_prefix(item) for item in SECTIONS[1][1]],
    ),
    (
        "2. Materials and methods",
        MATERIALS_AND_METHODS_REWRITTEN,
    ),
    (
        "3. Results and discussion",
        RESULTS_AND_DISCUSSION_REWRITTEN,
    ),
    (
        "4. Conclusions",
        CONCLUSIONS_REWRITTEN,
    ),
    SECTIONS[7],
    SECTIONS[8],
    SECTIONS[9],
]


TABLES = {
    "Table 1. Available raw-table variable ranges in the source Excel file.": [
        ["Variable", "Minimum", "Maximum", "Mean"],
        ["Stack total voltage", "78.9000", "202.0000", "164.4187"],
        ["Stack total current", "0.0000", "301.6000", "207.3628"],
        ["Stack power", "0.0000", "48.7782", "33.6780"],
        ["Single-cell voltage columns (global)", "0.0000", "0.7230", "0.1932"],
        ["Total impedance", "31.7300", "134.9300", "46.3499"],
        ["Mean impedance", "0.1300", "0.5840", "0.1927"],
        ["Maximum impedance", "0.1990", "0.7230", "0.2961"],
        ["Second-highest impedance", "0.1960", "0.7000", "0.2847"],
        ["Minimum impedance", "0.0650", "0.3090", "0.1244"],
        ["Second-lowest impedance", "0.0710", "0.3110", "0.1289"],
        ["Standard deviation", "0.0210", "0.1120", "0.0354"],
        ["EIS resistance real part", "0.0000", "0.3010", "0.0565"],
        ["EIS resistance imaginary part", "-0.0590", "0.0190", "0.0134"],
    ],
    "Table 2. Clean-test performance of the proposed model under the evaluated seed-44 fixed 8:2 protocol.": [
        ["Model", "Accuracy (%)", "Macro-F1 (%)", "Weighted-F1 (%)", "Class-0 P/R/F1 (%)", "Test windows", "Params"],
        ["Official-CAPT-UniShape-RBF-KANFusion", "100.00", "100.00", "100.00", "100.00 / 100.00 / 100.00", "55", "6,489,373"],
    ],
    "Table 3. Available aligned 8:2 clean baseline results from project artifacts.": [
        ["Model", "Category", "Accuracy (%)", "Macro-F1 (%)", "Weighted-F1 (%)", "Params"],
        ["Logistic regression", "Traditional ML", "94.55", "88.31", "94.24", "2,151"],
        ["Random forest", "Traditional ML", "94.55", "88.31", "94.24", "2,436"],
        ["MLP", "Deep learning", "92.73", "83.76", "92.07", "50,243"],
        ["1D CNN", "Deep learning", "85.45", "63.33", "78.91", "18,947"],
        ["Transformer", "Transformer", "94.55", "93.29", "94.86", "109,763"],
        ["iTransformer", "Transformer", "92.73", "83.76", "92.07", "214,595"],
    ],
    "Table 4. Ablation study under seed-44 fixed 8:2 protocol.": [
        ["Variant", "Accuracy (%)", "Macro-F1 (%)", "Weighted-F1 (%)", "Class-0 recall (%)", "Class-0 F1 (%)"],
        ["full_rbf", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["no_rbf", "98.18", "96.33", "98.16", "87.50", "93.33"],
        ["no_kan_fusion", "96.36", "92.46", "96.26", "75.00", "85.71"],
        ["static_prototype", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["no_transport_reg", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["no_separation_reg", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["no_eis_input", "87.27", "82.99", "87.56", "62.50", "58.82"],
        ["no_condition_input", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["stack_only", "83.64", "79.95", "84.27", "75.00", "85.71"],
        ["eis_cond_only", "100.00", "100.00", "100.00", "100.00", "100.00"],
    ],
    "Table 5. Proposed-model performance under joint-input test noise.": [
        ["SNR", "Accuracy (%)", "Macro-F1 (%)", "Weighted-F1 (%)", "Class-0 precision (%)", "Class-0 recall (%)", "Class-0 F1 (%)"],
        ["Clean", "100.00", "100.00", "100.00", "100.00", "100.00", "100.00"],
        ["30 dB", "85.45", "62.26", "78.79", "0.00", "0.00", "0.00"],
        ["20 dB", "87.27", "68.32", "82.84", "100.00", "12.50", "22.22"],
        ["10 dB", "83.64", "73.14", "83.26", "42.86", "37.50", "40.00"],
        ["5 dB", "60.00", "52.38", "62.86", "15.38", "25.00", "19.05"],
        ["0 dB", "25.45", "22.74", "14.82", "22.22", "25.00", "23.53"],
    ],
}


FIGURE_CAPTIONS = [
    "Fig. 1. Overall CAPT-UniShape architecture. The model receives x_op, constructed x_eis, and x_cond; extracts operation and EIS/impedance shape embeddings with two official UniShape backbones; encodes conditions with an MLP; fuses the three embeddings by Residual KAN-Fusion; and classifies with a condition-aware RBF prototype transport head.",
    "Fig. 2. Condition-aware RBF prototype transport head. The condition token is mapped through RBF bases to class-wise prototype offsets, which shift static prototypes before cosine-logit classification.",
    "Fig. 3. Clean-test confusion matrix for the proposed model under the evaluated seed-44 fixed 8:2 protocol. Class 2 is reported as Flooding / over-wet.",
    "Fig. 4. Ablation comparison of Accuracy and Macro-F1 under the evaluated seed-44 fixed 8:2 protocol.",
    "Fig. 5. Noise robustness curves for clean and SNR-based joint-input perturbation tests.",
]

FIGURES = [
    (ROOT / "outputs" / "paper_figures" / "architecture_diagram.png", FIGURE_CAPTIONS[0]),
    (ROOT / "outputs" / "paper_figures" / "prototype_head_diagram.png", FIGURE_CAPTIONS[1]),
    (ROOT / "outputs" / "paper_figures" / "confusion_matrix_proposed_seed44_8_2.png", FIGURE_CAPTIONS[2]),
    (ROOT / "outputs" / "paper_figures" / "ablation_bar_chart.png", FIGURE_CAPTIONS[3]),
    (ROOT / "outputs" / "paper_figures" / "noise_robustness_proposed_snr.png", FIGURE_CAPTIONS[4]),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 30 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.style = "Caption"
    p.add_run(caption).bold = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, text, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.6))
    else:
        doc.add_paragraph(f"[Figure file missing: {image_path}]")
    cap = doc.add_paragraph()
    cap.style = "Caption"
    cap.add_run(caption)
    doc.add_paragraph()


EQUATION_LINEAR_TEXT = {
    "(1)    x_i^eis = [c_i; g_i; a_i; q] in R^{4 x 128}.": r"x_i^{eis}=\left[c_i;g_i;a_i;q\right]\in\mathbb{R}^{4\times128}#(1)",
    "(2)    z_i^op = B_op(x_i^op),      z_i^eis = B_eis(x_i^eis),      z_i^op, z_i^eis in R^d.": r"z_i^{op}=B_{op}(x_i^{op}),\quad z_i^{eis}=B_{eis}(x_i^{eis}),\quad z_i^{op},z_i^{eis}\in\mathbb{R}^{d}#(2)",
    "(3)    z_i^cond = phi_cond(x_i^cond) = W_2 GELU(LN(W_1 x_i^cond + b_1)) + b_2.": r"z_i^{cond}=\phi_{cond}(x_i^{cond})=W_2GELU(LN(W_1x_i^{cond}+b_1))+b_2#(3)",
    "(4)    h_i = MLP(u_i) + lambda_KAN W_K KAN(Bottleneck(u_i)).": r"h_i=MLP(u_i)+\lambda_{KAN}W_KKAN(Bottleneck(u_i))#(4)",
    "(5)    p_{i,k} = p_k^0 + Delta p_{i,k}(z_i^cond).": r"p_{i,k}=p_k^0+\Delta p_{i,k}(z_i^{cond})#(5)",
    "(6)    r_{i,j} = exp(-||z_i^cond - c_j||_2^2 / (2 sigma_j^2)).": r"r_{i,j}=\exp\left(-\frac{\left\|z_i^{cond}-c_j\right\|_2^2}{2\sigma_j^2}\right)#(6)",
    "(7)    logit_{i,k} = cos(h_i, p_{i,k}) / tau.": r"logit_{i,k}=\frac{\cos(h_i,p_{i,k})}{\tau}#(7)",
    "(8)    L = L_CE + alpha_transport L_transport + alpha_sep L_sep + alpha_KAN L_KAN.": r"L=L_{CE}+\alpha_{transport}L_{transport}+\alpha_{sep}L_{sep}+\alpha_{KAN}L_{KAN}#(8)",
    "(9)    L_transport = mean_{i,k} ||Delta p_{i,k}||_2^2,": r"L_{transport}=mean_{i,k}\left\|\Delta p_{i,k}\right\|_2^2#(9)",
    "(10)    L_sep = mean_{a != b} max(0, cos(p_a^0, p_b^0) - m).": r"L_{sep}=mean_{a\ne b}\max(0,\cos(p_a^0,p_b^0)-m)#(10)",
}


def add_equation(doc: Document, linear_text: str) -> None:
    """Add a Word OMML equation paragraph using Word linear equation syntax."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    omath_para = OxmlElement("m:oMathPara")
    omath_para_pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc")
    jc.set(qn("m:val"), "center")
    omath_para_pr.append(jc)
    omath_para.append(omath_para_pr)

    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    run_props = OxmlElement("m:rPr")
    sty = OxmlElement("m:sty")
    sty.set(qn("m:val"), "p")
    run_props.append(sty)
    math_run.append(run_props)
    text = OxmlElement("m:t")
    text.text = linear_text
    math_run.append(text)
    omath.append(math_run)
    omath_para.append(omath)
    p._p.append(omath_para)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Caption", 10)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Caption"].font.italic = True


def add_front_matter(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = "Title"
    run = title.add_run(TITLE)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author names and affiliations: to be completed").italic = True

    doc.add_heading("Highlights", level=1)
    for item in HIGHLIGHTS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT)

    doc.add_heading("Keywords", level=1)
    doc.add_paragraph("; ".join(KEYWORDS))

    doc.add_heading("Abbreviations", level=1)
    add_table(doc, "Abbreviations used in this manuscript.", ABBREVIATIONS)


def add_sections(doc: Document) -> None:
    table_inserted = False
    for heading, paragraphs in JOURNAL_SECTIONS:
        doc.add_heading(heading, level=1)
        for para in paragraphs:
            if para in EQUATION_LINEAR_TEXT:
                add_equation(doc, EQUATION_LINEAR_TEXT[para])
            else:
                doc.add_paragraph(para)

        if heading == "2. Materials and methods":
            for caption in [
                "Table 1. Available raw-table variable ranges in the source Excel file.",
                "Table 2. Clean-test performance of the proposed model under the evaluated seed-44 fixed 8:2 protocol.",
                "Table 3. Available aligned 8:2 clean baseline results from project artifacts.",
            ]:
                add_table(doc, caption, TABLES[caption])
            for image_path, caption in FIGURES[:2]:
                add_figure(doc, image_path, caption)
        if heading == "3. Results and discussion":
            for caption in [
                "Table 4. Ablation study under seed-44 fixed 8:2 protocol.",
                "Table 5. Proposed-model performance under joint-input test noise.",
            ]:
                add_table(doc, caption, TABLES[caption])
            for image_path, caption in FIGURES[2:]:
                add_figure(doc, image_path, caption)
            table_inserted = True
    assert table_inserted


def write_markdown() -> None:
    lines: list[str] = [f"# {TITLE}", "", "## Highlights"]
    lines.extend([f"- {item}" for item in HIGHLIGHTS])
    lines.extend(["", "## Abstract", ABSTRACT, "", "## Keywords", "; ".join(KEYWORDS), ""])
    lines.extend(["", "## Abbreviations", ""])
    lines.append("| " + " | ".join(ABBREVIATIONS[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(ABBREVIATIONS[0])) + " |")
    for row in ABBREVIATIONS[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")
    for heading, paragraphs in JOURNAL_SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        for para in paragraphs:
            lines.append(EQUATION_LINEAR_TEXT.get(para, para))
        lines.append("")
        if heading == "2. Materials and methods":
            for caption in [
                "Table 1. Available raw-table variable ranges in the source Excel file.",
                "Table 2. Clean-test performance of the proposed model under the evaluated seed-44 fixed 8:2 protocol.",
                "Table 3. Available aligned 8:2 clean baseline results from project artifacts.",
            ]:
                lines.append(caption)
                rows = TABLES[caption]
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
            for image_path, caption in FIGURES[:2]:
                rel = image_path.relative_to(ROOT).as_posix()
                lines.extend([f"![{caption}]({rel})", caption, ""])
        if heading == "3. Results and discussion":
            for caption in [
                "Table 4. Ablation study under seed-44 fixed 8:2 protocol.",
                "Table 5. Proposed-model performance under joint-input test noise.",
            ]:
                lines.append(caption)
                rows = TABLES[caption]
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                for row in rows[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
            for image_path, caption in FIGURES[2:]:
                rel = image_path.relative_to(ROOT).as_posix()
                lines.extend([f"![{caption}]({rel})", caption, ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    add_front_matter(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_sections(doc)
    doc.save(DOCX_PATH)
    write_markdown()


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)
