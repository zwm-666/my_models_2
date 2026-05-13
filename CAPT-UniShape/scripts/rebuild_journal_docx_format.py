from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "CAPT-UniShape_Elsevier_revised_draft.docx"
OUT = ROOT / "docs" / "CAPT-UniShape_Elsevier_revised_draft_reformatted.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W, "m": M}


EQUATIONS = {
    1: "x_i^(eis)=[c_i;g_i;a_i;q]∈ℝ^(4×128)#(1)",
    2: "z_i^(op)=B_(op)(x_i^(op)), z_i^(eis)=B_(eis)(x_i^(eis)), z_i^(op),z_i^(eis)∈ℝ^d#(2)",
    3: "z_i^(cond)=φ_(cond)(x_i^(cond))=W_2 GELU(LN(W_1 x_i^(cond)+b_1))+b_2#(3)",
    4: "h_i=MLP(u_i)+λ_(KAN) W_K KAN(Bottleneck(u_i))#(4)",
    5: "p_(i,k)=p_k^0+Δp_(i,k)(z_i^(cond))#(5)",
    6: "r_(i,j)=exp(-(∥z_i^(cond)-c_j∥_2^2)/(2σ_j^2))#(6)",
    7: "logit_(i,k)=cos(h_i,p_(i,k))/τ#(7)",
    8: "L=L_(CE)+α_(transport) L_(transport)+α_(sep) L_(sep)+α_(KAN) L_(KAN)#(8)",
    9: "L_(transport)=mean_(i,k) ∥Δp_(i,k)∥_2^2#(9)",
    10: "L_(sep)=mean_(a≠b) max(0,cos(p_a^0,p_b^0)-m)#(10)",
}


MATH_TOKENS = {
    "[N, 3, 64]": "[N,3,64]",
    "[N, 4, 128]": "[N,4,128]",
    "[N, 12]": "[N,12]",
    "R^{N x 3 x 64}": "ℝ^(N×3×64)",
    "R^{N x 4 x 128}": "ℝ^(N×4×128)",
    "R^{N x 12}": "ℝ^(N×12)",
    "R^{3 x 64}": "ℝ^(3×64)",
    "R^{4 x 128}": "ℝ^(4×128)",
    "R^{12}": "ℝ^12",
    "R^{3d}": "ℝ^(3d)",
    "R^{d}": "ℝ^d",
    "R^9": "ℝ^9",
    "R^d": "ℝ^d",
    "y_hat_i = argmax_k f_theta(x_i)_k": "ŷ_i=argmax_k f_θ(x_i)_k",
    "q in [0, 1]": "q∈[0,1]",
    "argmax_k": "argmax_k",
    "f_theta": "f_θ",
    "y_hat_i": "ŷ_i",
    "Delta p_{i,k}": "Δp_(i,k)",
    "Delta p": "Δp",
    "alpha_transport": "α_(transport)",
    "alpha_sep": "α_(sep)",
    "alpha_KAN": "α_(KAN)",
    "lambda_KAN": "λ_(KAN)",
    "phi_cond": "φ_(cond)",
    "sigma_j": "σ_j",
    "tau": "τ",
    "x_i^cond": "x_i^(cond)",
    "x_i^eis": "x_i^(eis)",
    "x_i^op": "x_i^(op)",
    "z_i^cond": "z_i^(cond)",
    "z_i^eis": "z_i^(eis)",
    "z_i^op": "z_i^(op)",
    "p_{i,k}": "p_(i,k)",
    "r_{i,j}": "r_(i,j)",
    "logit_{i,k}": "logit_(i,k)",
    "mean_{i,k}": "mean_(i,k)",
    "mean_{a != b}": "mean_(a≠b)",
    "p_a^0": "p_a^0",
    "p_b^0": "p_b^0",
    "p_k^0": "p_k^0",
    "s_i": "s_i",
    "x_i": "x_i",
    "c_i": "c_i",
    "g_i": "g_i",
    "a_i": "a_i",
    "c_j": "c_j",
    "h_i": "h_i",
    "u_i": "u_i",
    "y_i": "y_i",
    "B_op": "B_(op)",
    "B_eis": "B_(eis)",
    "W_1": "W_1",
    "W_2": "W_2",
    "W_K": "W_K",
    "b_1": "b_1",
    "b_2": "b_2",
    "L_transport": "L_(transport)",
    "L_CE": "L_(CE)",
    "L_sep": "L_(sep)",
    "L_KAN": "L_(KAN)",
    "x_cond": "x_(cond)",
    "x_eis": "x_(eis)",
    "x_op": "x_(op)",
    "z_cond": "z_(cond)",
    "z_eis": "z_(eis)",
    "z_op": "z_(op)",
    "logit_k": "logit_k",
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(doc: Document, name: str, size: float, bold: bool = False, italic: bool = False) -> None:
    style = doc.styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)


def ensure_para_style(doc: Document, name: str, base: str = "Normal"):
    styles = doc.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, 1)


def apply_docx_styles(path: Path) -> None:
    doc = Document(path)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    set_style(doc, "Normal", 11)
    set_style(doc, "Title", 16, bold=True)
    set_style(doc, "Heading 1", 12, bold=True)
    set_style(doc, "Caption", 10)
    if "List Bullet" in doc.styles:
        set_style(doc, "List Bullet", 11)

    eq_style = ensure_para_style(doc, "Equation")
    eq_style.font.name = "Times New Roman"
    eq_style.font.size = Pt(11)

    normal = doc.styles["Normal"].paragraph_format
    normal.first_line_indent = Pt(24)
    normal.line_spacing = 1.15
    normal.space_after = Pt(6)
    normal.space_before = Pt(0)

    h1 = doc.styles["Heading 1"].paragraph_format
    h1.first_line_indent = Pt(0)
    h1.space_before = Pt(12)
    h1.space_after = Pt(6)
    h1.keep_with_next = True

    title = doc.styles["Title"].paragraph_format
    title.first_line_indent = Pt(0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    cap = doc.styles["Caption"].paragraph_format
    cap.first_line_indent = Pt(0)
    cap.line_spacing = 1.0
    cap.space_before = Pt(6)
    cap.space_after = Pt(6)
    cap.keep_with_next = True

    eq_fmt = eq_style.paragraph_format
    eq_fmt.first_line_indent = Pt(0)
    eq_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_fmt.space_before = Pt(6)
    eq_fmt.space_after = Pt(6)

    in_references = False
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        has_math = "<m:oMath" in para._p.xml

        if idx == 0:
            para.style = doc.styles["Title"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text in {"Declaration of competing interest", "Data availability", "References"}:
            para.style = doc.styles["Heading 1"]
            para.paragraph_format.first_line_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if text == "References":
                in_references = True
        elif text == "Author names and affiliations: to be completed":
            para.paragraph_format.first_line_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif para.style.name == "Heading 1":
            para.paragraph_format.first_line_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif para.style.name == "Caption" or text.startswith(("Fig. ", "Table ")) or text == "Abbreviations used in this manuscript.":
            para.style = doc.styles["Caption"]
            para.paragraph_format.first_line_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif has_math and not text:
            para.style = eq_style
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
        elif para.style.name.startswith("List"):
            para.paragraph_format.first_line_indent = Pt(0)
        elif text.startswith("Proton exchange membrane fuel cell;"):
            para.paragraph_format.first_line_indent = Pt(0)
        elif in_references and text.startswith("["):
            para.paragraph_format.first_line_indent = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.paragraph_format.first_line_indent = Pt(24)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in para.runs:
            if para.style.name == "Title":
                set_run_font(run, 16, True)
            elif para.style.name == "Heading 1":
                set_run_font(run, 12, True)
            elif para.style.name == "Caption":
                set_run_font(run, 10)
            else:
                set_run_font(run, 11)

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        set_run_font(run, 9, bold=(row_idx == 0))

    doc.save(path)


def w_text(text: str):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(fonts)
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def m_run(text: str):
    om = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = text
    mr.append(mt)
    om.append(mr)
    return om


def display_math(text: str):
    para = OxmlElement("m:oMathPara")
    om = m_run(text)
    para.append(om)
    return para


def replace_equation_paragraphs_and_inline_math(path: Path) -> None:
    import lxml.etree as ET

    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        xml = zin.read("word/document.xml")
        root = ET.fromstring(xml)
        body = root.find("w:body", NS)

        eq_counter = 0
        def token_pattern(token: str) -> str:
            pattern = re.escape(token)
            if re.match(r"[A-Za-z0-9_]", token[0]):
                pattern = r"(?<![A-Za-z0-9_])" + pattern
            if re.match(r"[A-Za-z0-9_]", token[-1]):
                pattern = pattern + r"(?![A-Za-z0-9_])"
            return pattern

        token_re = re.compile(
            "|".join(token_pattern(k) for k in sorted(MATH_TOKENS, key=len, reverse=True))
        )

        for p in body.findall(".//w:p", NS):
            text = "".join(t.text or "" for t in p.findall(".//w:t", NS)).strip()
            has_math = p.find(".//m:oMath", NS) is not None or p.find(".//m:oMathPara", NS) is not None

            if has_math and not text:
                eq_counter += 1
                ppr = p.find("w:pPr", NS)
                for child in list(p):
                    if child is not ppr:
                        p.remove(child)
                p.append(display_math(EQUATIONS.get(eq_counter, "")))
                continue

            if not text or has_math:
                continue

            full_text = "".join(t.text or "" for t in p.findall(".//w:t", NS))
            if not token_re.search(full_text):
                continue

            ppr = p.find("w:pPr", NS)
            for child in list(p):
                if child is not ppr:
                    p.remove(child)

            pos = 0
            for match in token_re.finditer(full_text):
                if match.start() > pos:
                    p.append(w_text(full_text[pos : match.start()]))
                token = match.group(0)
                p.append(m_run(MATH_TOKENS[token]))
                pos = match.end()
            if pos < len(full_text):
                p.append(w_text(full_text[pos:]))

        zout.writestr("word/document.xml", ET.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True))
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                continue
            zout.writestr(item, zin.read(item.filename))
    tmp.replace(path)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    apply_docx_styles(OUT)
    replace_equation_paragraphs_and_inline_math(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
